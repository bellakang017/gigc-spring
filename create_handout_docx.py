#!/usr/bin/env python3
"""Create a professional Word document version of the IWG Handout."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Helvetica Neue'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x42, 0x42, 0x45)

pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.4

# Colors
NOIR = RGBColor(0x1D, 0x1D, 0x1F)
CHARCOAL = RGBColor(0x42, 0x42, 0x45)
GREY = RGBColor(0x6E, 0x6E, 0x73)
SILVER = RGBColor(0x86, 0x86, 0x8B)
GOLD = RGBColor(0xC9, 0xA9, 0x6E)
GOLD_DEEP = RGBColor(0xA6, 0x8B, 0x4B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_gold_line(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('___________')
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_section_title(doc, text, accent_word=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    if accent_word and accent_word in text:
        parts = text.split(accent_word)
        if parts[0]:
            run = p.add_run(parts[0])
            run.font.size = Pt(18)
            run.font.color.rgb = NOIR
        run = p.add_run(accent_word)
        run.font.size = Pt(18)
        run.font.color.rgb = GOLD
        run.italic = True
        if len(parts) > 1 and parts[1]:
            run = p.add_run(parts[1])
            run.font.size = Pt(18)
            run.font.color.rgb = NOIR
    else:
        run = p.add_run(text)
        run.font.size = Pt(18)
        run.font.color.rgb = NOIR

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(14)

def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', 'D2D2D7'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)

# ═══════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════

# Badge
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
run = p.add_run('International Job Search Working Group  ·  Workshop #2')
run.font.size = Pt(8)
run.font.color.rgb = CHARCOAL
run.font.bold = True

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Preparing for ')
run.font.size = Pt(26)
run.font.color.rgb = NOIR
run = p.add_run('AI Collaboration')
run.font.size = Pt(26)
run.font.color.rgb = GOLD
run.italic = True
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run('Interviews')
run.font.size = Pt(26)
run.font.color.rgb = NOIR

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Companies are adding AI to their interview process.\nIf that sounds intimidating — don\'t worry,\nwe\'ll break it down together.')
run.font.size = Pt(9.5)
run.font.color.rgb = GREY
run.italic = True

# Meta
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Tuesday, February 24, 2026  ·  3:00–4:00 PM  ·  FAC 2.134')
run.font.size = Pt(8)
run.font.color.rgb = SILVER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Office of Career & Life Design  ·  UT Austin')
run.font.size = Pt(8)
run.font.color.rgb = SILVER

add_gold_line(doc)

# ═══════════════════════════════════════════
# WHY THIS MATTERS
# ═══════════════════════════════════════════

add_section_title(doc, 'What\'s Changing', 'Changing')
add_subtitle(doc, 'The job market is shifting fast. Here\'s what\'s happening right now.')

# Stats table
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

stats = [
    ('93%', 'of recruiters plan to increase\nAI in hiring this year'),
    ('77%', 'of hiring teams regularly encounter\nAI-generated applications'),
    ('$2.7B', 'BCG revenue from\nAI consulting last year'),
]

for i, (num, label) in enumerate(stats):
    cell = table.cell(0, i)
    set_cell_shading(cell, 'F5F5F7')
    border_style = {'val': 'single', 'sz': '4', 'color': 'D2D2D7'}
    set_cell_border(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(num)
    run.font.size = Pt(22)
    run.font.color.rgb = NOIR
    run.font.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    run = p2.add_run(label)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY

# Insight box
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(14)
run = p.add_run('The good news: ')
run.font.bold = True
run.font.color.rgb = NOIR
run.font.size = Pt(9.5)
run = p.add_run('The top skills employers want are critical thinking, problem solving, time management, and adaptability — skills AI can\'t replicate. What companies are really testing isn\'t technical AI expertise. It\'s ')
run.font.size = Pt(9.5)
run.font.color.rgb = CHARCOAL
run = p.add_run('structured thinking, clear communication, and sound judgment.')
run.font.bold = True
run.font.color.rgb = NOIR
run.font.size = Pt(9.5)

add_gold_line(doc)

# ═══════════════════════════════════════════
# WHAT YOU'LL DO
# ═══════════════════════════════════════════

add_section_title(doc, 'What You\'ll Practice', 'Practice')
add_subtitle(doc, 'Four frameworks that transfer to any interview, any field.')

frameworks = [
    ('1', 'ASK → Frame Clear Questions', 'Turn vague problems into specific, answerable questions — the exact skill companies test when candidates work with AI tools.'),
    ('2', 'STRUCTURE → Break It Down', 'Master MECE thinking: split any messy problem into clean buckets that cover everything. A golden rule across consulting and tech.'),
    ('3', 'REFINE → Review & Iterate', 'Practice evaluating information critically and improving your first response under time pressure. Use AI as a support tool while you drive.'),
    ('4', 'STORY → Experience → Impact', 'Transform "I worked on a research project" into "I identified a gap, designed a methodology, and shifted how the field understands X."'),
]

table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for idx, (num, title, desc) in enumerate(frameworks):
    row = idx // 2
    col = idx % 2
    cell = table.cell(row, col)
    set_cell_shading(cell, 'FAFAFA')
    border_style = {'val': 'single', 'sz': '4', 'color': 'D2D2D7'}
    set_cell_border(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(num + '  ')
    run.font.size = Pt(14)
    run.font.color.rgb = GOLD
    run.font.bold = True

    run = p.add_run(title)
    run.font.size = Pt(10)
    run.font.color.rgb = NOIR
    run.font.bold = True

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    run = p2.add_run(desc)
    run.font.size = Pt(8.5)
    run.font.color.rgb = GREY

add_gold_line(doc)

# ═══════════════════════════════════════════
# INTERNATIONAL SCHOLARS
# ═══════════════════════════════════════════

add_section_title(doc, 'Why This Matters for International Scholars', 'International Scholars')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run = p.add_run('Your global perspective, cross-cultural fluency, and adaptability are exactly what companies need for AI-era challenges. But you have to know how to communicate that strategically. This workshop gives you the frameworks to translate your unique experiences into competitive advantage.')
run.font.size = Pt(9.5)
run.font.color.rgb = CHARCOAL

# Stats
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
intl_stats = [
    ('50%', 'of employers say applicants lack relevant experience'),
    ('26%', 'struggle to evaluate informal or self-taught skills'),
]
for i, (num, label) in enumerate(intl_stats):
    cell = table.cell(0, i)
    set_cell_shading(cell, '1D1D1F')
    border_style = {'val': 'single', 'sz': '4', 'color': '1D1D1F'}
    set_cell_border(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(num)
    run.font.size = Pt(18)
    run.font.color.rgb = GOLD
    run.font.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    run = p2.add_run(label)
    run.font.size = Pt(8)
    run.font.color.rgb = WHITE

add_gold_line(doc)

# ═══════════════════════════════════════════
# SKILLS TRANSFER
# ═══════════════════════════════════════════

add_section_title(doc, 'These Skills Transfer Everywhere', 'Transfer')

transfers = [
    'Technical interviews at any company',
    'Research proposal defenses & grants',
    'Networking conversations & elevator pitches',
    'Job talks explaining research to non-experts',
]

table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for idx, text in enumerate(transfers):
    row = idx // 2
    col = idx % 2
    cell = table.cell(row, col)
    set_cell_shading(cell, 'FAFAFA')
    border_style = {'val': 'single', 'sz': '4', 'color': 'D2D2D7'}
    set_cell_border(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('  ✓  ')
    run.font.color.rgb = GOLD
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = CHARCOAL

# ═══════════════════════════════════════════
# PAGE BREAK → TIMELINE
# ═══════════════════════════════════════════

doc.add_page_break()

add_section_title(doc, 'Workshop Timeline', 'Timeline')
add_subtitle(doc, '50 minutes of structured practice.')

timeline = [
    ('3:00', 'Welcome & Name Tags', 'Settle in, find your table, pick up your name tag and handout.'),
    ('3:05', 'Find Your Instinct', 'Discover your natural interview style. Pick your animal, identify your superpower.'),
    ('3:12', 'Build Your HIVE Story', 'Use your story prompt: Habitat → Initiative → Venture → Effect.'),
    ('3:22', '60-Second Partner Practice', 'Tell your HIVE story in 60 seconds. Partner gives feedback. Switch.'),
    ('3:32', 'Sharpen Your Instinct', 'Practice thinking frameworks: ASK → STRUCTURE → REFINE.'),
    ('3:48', 'Wrap-up', 'Key takeaways and how to keep practicing.'),
]

table = doc.add_table(rows=len(timeline), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (time, title, desc) in enumerate(timeline):
    cell_time = table.cell(i, 0)
    cell_content = table.cell(i, 1)

    set_cell_shading(cell_time, 'FAFAFA')
    border_style = {'val': 'single', 'sz': '4', 'color': 'D2D2D7'}
    set_cell_border(cell_time, top=border_style, bottom=border_style, left=border_style, right=border_style)
    set_cell_border(cell_content, top=border_style, bottom=border_style, left=border_style, right=border_style)

    p = cell_time.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(time)
    run.font.size = Pt(10)
    run.font.color.rgb = GOLD
    run.font.bold = True

    p = cell_content.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(title)
    run.font.size = Pt(10)
    run.font.color.rgb = NOIR
    run.font.bold = True

    p2 = cell_content.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    run = p2.add_run(desc)
    run.font.size = Pt(8.5)
    run.font.color.rgb = GREY

# Set time column width
for row in table.rows:
    row.cells[0].width = Cm(2.5)

add_gold_line(doc)

# ═══════════════════════════════════════════
# SITE VISITS
# ═══════════════════════════════════════════

# ═══════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Bella Kang')
run.font.size = Pt(10)
run.font.color.rgb = NOIR
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('M.A. Advertising  ·  Graduate Global Impact Consultant')
run.font.size = Pt(8)
run.font.color.rgb = SILVER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Office of Career & Life Design, UT Austin  ·  Spring 2026')
run.font.size = Pt(8)
run.font.color.rgb = SILVER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('all scholars welcome ✦')
run.font.size = Pt(11)
run.font.color.rgb = GOLD
run.italic = True

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'IWG_Workshop2_Handout.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
